from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "doc" / "generated"
OUT_FILE = OUT_DIR / "ElderGo_Route_Scoring_Formula_Insert.docx"

NAVY = "1E3A5F"
BLUE = "4A90E2"
GRAY = "374151"


def style_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.12

    for name, size, color in [
        ("Title", 20, NAVY),
        ("Heading 1", 16, NAVY),
        ("Heading 2", 13, BLUE),
        ("Heading 3", 11.5, NAVY),
    ]:
        style = styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)


def add_formula(doc: Document, text: str, size: float = 15) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run(text)
    run.italic = True
    run.font.name = "Cambria Math"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Cambria Math")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(GRAY)


def add_bullet(doc: Document, label: str, body: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    r1 = p.add_run(label)
    r1.bold = True
    r1.font.color.rgb = RGBColor.from_string(NAVY)
    p.add_run(body)


def build_docx() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    style_doc(doc)

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("Route Scoring Formula Insert")

    note = doc.add_paragraph()
    note.paragraph_format.space_after = Pt(12)
    r = note.add_run("Where to paste: ")
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(NAVY)
    note.add_run(
        "Place this under Section 5.2 Backend Route Processing Flow or under the new "
        "Route Recommendation Decision Logic subsection."
    )

    doc.add_heading("Route Scoring Formula", level=1)
    doc.add_paragraph(
        "The current ElderGo KL route recommendation logic evaluates candidate routes returned by Google Directions "
        "using a cost-based scoring model. Each candidate route is compared using total travel duration, walking "
        "distance, number of transfers, and the user's selected travel preferences. The system selects the route "
        "with the lowest final cost."
    )

    add_formula(doc, "Cost = M + (Ww x D) + (Wt x T) + (Wa x A)")

    doc.add_paragraph("Where:")
    add_bullet(doc, "M ", "represents the total route duration in minutes.")
    add_bullet(doc, "D ", "represents the total walking distance in metres.")
    add_bullet(doc, "T ", "represents the number of transfers.")
    add_bullet(doc, "A ", "represents the unknown accessibility risk. In the current implementation, this value is fixed as 1.")
    add_bullet(doc, "Ww ", "represents the walking-distance weight.")
    add_bullet(doc, "Wt ", "represents the transfer-count weight.")
    add_bullet(doc, "Wa ", "represents the accessibility-risk weight.")

    doc.add_heading("Current Iteration 2 Weight Values", level=2)
    doc.add_paragraph("The current backend implementation applies the following weights:")

    add_formula(doc, "Ww = 0.04 if Least Walk is enabled; otherwise 0.02", 13)
    add_formula(doc, "Wt = 14 if Fewest Transfers is enabled; otherwise 8", 13)
    add_formula(doc, "Wa = 50 if Accessibility First is enabled; otherwise 15", 13)

    doc.add_heading("Backend Formula in Current Build", level=2)
    doc.add_paragraph("Therefore, the current backend scoring formula can be expressed as:")
    add_formula(
        doc,
        "Cost = DurationMinutes + (WalkingDistanceMeters x Ww) + (Transfers x Wt) + (1 x Wa)",
        13.5,
    )

    doc.add_paragraph(
        "A lower cost represents a more suitable route. ElderGo KL selects the candidate route with the lowest "
        "cost value. When Accessibility First is enabled, the backend first prioritises candidates with lower "
        "walking distance and fewer transfers before using the weighted cost as a tie-breaker. This reflects the "
        "current implementation, where lower walking burden and fewer transfers are used as practical proxies for "
        "elderly-friendly accessibility."
    )

    doc.add_heading("Weather Boundary", level=2)
    doc.add_paragraph(
        "Weather is not included in the current scoring formula. In Iteration 2, weather is displayed as advisory "
        "information after the route recommendation is generated and does not automatically change or replace the "
        "selected route."
    )

    doc.save(OUT_FILE)


if __name__ == "__main__":
    build_docx()
