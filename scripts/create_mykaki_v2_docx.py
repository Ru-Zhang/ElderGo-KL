from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "doc" / "generated"
DIAGRAM_DIR = OUT_DIR / "diagrams"
OUT_FILE = OUT_DIR / "MyKaki_System_Architecture_and_Technical_Design_Specification_v2.0.docx"

NAVY = "1E3A5F"
BLUE = "4A90E2"
GREEN = "6BBF59"
ORANGE = "E67E22"
LIGHT_BLUE = "EAF3FE"
LIGHT_GREEN = "ECF7EA"
LIGHT_ORANGE = "FFF3E8"
LIGHT_GRAY = "F5F7FA"
BORDER = "D9E2EC"
WHITE = "FFFFFF"


def _font(size: int, bold: bool = False):
    candidates = [
        Path("C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf"),
        Path("C:/Windows/Fonts/calibrib.ttf" if bold else "C:/Windows/Fonts/calibri.ttf"),
    ]
    for path in candidates:
        if path.exists():
            return ImageFont.truetype(str(path), size)
    return ImageFont.load_default()


def _wrap(draw: ImageDraw.ImageDraw, text: str, font, width: int) -> list[str]:
    words = text.split()
    lines: list[str] = []
    current = ""
    for word in words:
        candidate = f"{current} {word}".strip()
        if draw.textbbox((0, 0), candidate, font=font)[2] <= width:
            current = candidate
        else:
            if current:
                lines.append(current)
            current = word
    if current:
        lines.append(current)
    return lines


def _draw_box(draw, xy, text, fill="#EFEAFF", outline="#A88CFF", font_size=22, title=False):
    x1, y1, x2, y2 = xy
    draw.rounded_rectangle(xy, radius=10, fill=fill, outline=outline, width=3)
    font = _font(font_size, bold=title)
    lines = _wrap(draw, text, font, max(20, x2 - x1 - 28))
    line_h = font_size + 7
    total_h = len(lines) * line_h
    y = y1 + (y2 - y1 - total_h) / 2
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=font)
        draw.text((x1 + (x2 - x1 - (bbox[2] - bbox[0])) / 2, y), line, fill="#1E3A5F", font=font)
        y += line_h


def _arrow(draw, start, end, color="#334155", width=3):
    draw.line([start, end], fill=color, width=width)
    x1, y1 = start
    x2, y2 = end
    if abs(y2 - y1) >= abs(x2 - x1):
        direction = 1 if y2 >= y1 else -1
        pts = [(x2, y2), (x2 - 9, y2 - direction * 15), (x2 + 9, y2 - direction * 15)]
    else:
        direction = 1 if x2 >= x1 else -1
        pts = [(x2, y2), (x2 - direction * 15, y2 - 9), (x2 - direction * 15, y2 + 9)]
    draw.polygon(pts, fill=color)


def _caption(draw, text, x, y):
    draw.text((x, y), text, fill="#64748B", font=_font(20))


def save_architecture_diagram(path: Path) -> None:
    img = Image.new("RGB", (1500, 850), "white")
    d = ImageDraw.Draw(img)
    d.text((70, 45), "Figure 1: System Component Interaction Diagram (Iteration 2)", fill="#111827", font=_font(34, True))
    boxes = {
        "user": (70, 360, 280, 455, "Older Adult User\nMobile Browser"),
        "frontend": (380, 340, 610, 475, "ElderGo KL Frontend\nReact + Vite\nLocal Storage"),
        "api": (760, 325, 1010, 490, "ElderGo KL API\nFastAPI\nGunicorn/Uvicorn"),
        "pg": (1180, 135, 1435, 235, "PostgreSQL + PostGIS\nRail / Station /\nAccessibility Data"),
        "google": (1180, 285, 1435, 385, "Google APIs\nPlaces + Maps\nDirections"),
        "weather": (1180, 435, 1435, 535, "OpenWeather\nForecast API"),
        "render": (1180, 585, 1435, 685, "Render Services\nFrontend Static\nBackend Web"),
    }
    for key, (x1, y1, x2, y2, label) in boxes.items():
        fill = "#EAF3FE" if key in {"frontend", "api"} else "#F1ECFF"
        if key == "pg":
            fill = "#ECF7EA"
        if key == "weather":
            fill = "#FFF3E8"
        _draw_box(d, (x1, y1, x2, y2), label, fill=fill, font_size=22, title=key in {"frontend", "api"})
    _arrow(d, (280, 408), (380, 408))
    _caption(d, "uses app", 300, 375)
    _arrow(d, (610, 390), (760, 390))
    _caption(d, "API calls", 635, 358)
    _arrow(d, (610, 455), (760, 455))
    _caption(d, "settings fallback", 618, 470)
    _arrow(d, (1010, 365), (1180, 185))
    _caption(d, "SQL / PostGIS", 1035, 245)
    _arrow(d, (1010, 405), (1180, 335))
    _caption(d, "Places / Directions", 1035, 348)
    _arrow(d, (1010, 445), (1180, 485))
    _caption(d, "Forecast", 1045, 462)
    _arrow(d, (610, 520), (1180, 635))
    _caption(d, "static frontend hosted on Render", 690, 560)
    _arrow(d, (1010, 545), (1180, 635))
    _caption(d, "backend web service on Render", 1040, 570)
    path.parent.mkdir(parents=True, exist_ok=True)
    img.save(path)


def save_vertical_flow(path: Path, title: str, steps: list[str]) -> None:
    width = 1000
    box_w = 460
    box_h = 84
    gap = 42
    top = 95
    height = top + len(steps) * box_h + (len(steps) - 1) * gap + 80
    img = Image.new("RGB", (width, height), "white")
    d = ImageDraw.Draw(img)
    d.text((60, 35), title, fill="#111827", font=_font(34, True))
    x1 = (width - box_w) // 2
    for i, step in enumerate(steps):
        y1 = top + i * (box_h + gap)
        _draw_box(d, (x1, y1, x1 + box_w, y1 + box_h), step, font_size=22)
        if i < len(steps) - 1:
            _arrow(d, (width // 2, y1 + box_h), (width // 2, y1 + box_h + gap - 8))
    path.parent.mkdir(parents=True, exist_ok=True)
    img.save(path)


def save_backend_flow(path: Path) -> None:
    img = Image.new("RGB", (1300, 1000), "white")
    d = ImageDraw.Draw(img)
    d.text((60, 35), "Figure 3: Backend Route Request Processing Logic", fill="#111827", font=_font(34, True))
    _draw_box(d, (515, 100, 785, 170), "Receive POST\n/routes/recommend")
    diamond = [(650, 230), (780, 350), (650, 470), (520, 350)]
    d.polygon(diamond, fill="#F1ECFF", outline="#A88CFF")
    d.line(diamond + [diamond[0]], fill="#A88CFF", width=3)
    for line, y in zip(["Origin and", "destination", "valid?"], [315, 345, 375]):
        bbox = d.textbbox((0, 0), line, font=_font(22))
        d.text((650 - (bbox[2] - bbox[0]) / 2, y), line, fill="#1E3A5F", font=_font(22))
    _arrow(d, (650, 170), (650, 230))
    _arrow(d, (520, 350), (305, 510))
    _caption(d, "No", 415, 385)
    _draw_box(d, (120, 510, 390, 590), "Return validation\nor provider error")
    _arrow(d, (780, 350), (910, 350))
    _caption(d, "Yes", 805, 315)
    _draw_box(d, (910, 305, 1190, 395), "Fetch Google Maps\ntransit candidates")
    _arrow(d, (1050, 395), (1050, 465))
    _draw_box(d, (910, 465, 1190, 555), "Score by preferences:\naccessibility, walking,\ntransfers, duration")
    _arrow(d, (1050, 555), (1050, 625))
    _draw_box(d, (910, 625, 1190, 715), "Annotate steps with\nstation/accessibility\ncontext")
    _arrow(d, (1050, 715), (1050, 785))
    _draw_box(d, (910, 785, 1190, 875), "Persist route or use\ndemo/ephemeral ID")
    _arrow(d, (910, 830), (650, 830))
    _draw_box(d, (450, 785, 720, 875), "Return one\nRecommendedRoute")
    path.parent.mkdir(parents=True, exist_ok=True)
    img.save(path)


def save_sequence_flow(path: Path, title: str, actors: list[str], messages: list[tuple[int, int, str, bool]]) -> None:
    width = 1400
    height = 650
    img = Image.new("RGB", (width, height), "white")
    d = ImageDraw.Draw(img)
    d.text((60, 35), title, fill="#111827", font=_font(34, True))
    xs = [180 + i * ((width - 360) // (len(actors) - 1)) for i in range(len(actors))]
    for x, actor in zip(xs, actors):
        _draw_box(d, (x - 95, 100, x + 95, 165), actor, font_size=22)
        d.line([(x, 165), (x, 560)], fill="#D8CCFF", width=2)
        _draw_box(d, (x - 95, 560, x + 95, 625), actor, font_size=22)
    y = 225
    for frm, to, text, dashed in messages:
        x1, x2 = xs[frm], xs[to]
        if dashed:
            # simple dashed line
            segments = 24
            for i in range(segments):
                if i % 2 == 0:
                    sx = x1 + (x2 - x1) * i / segments
                    ex = x1 + (x2 - x1) * (i + 1) / segments
                    d.line([(sx, y), (ex, y)], fill="#334155", width=2)
            _arrow(d, (x2 - (1 if x2 > x1 else -1) * 20, y), (x2, y), width=2)
        else:
            _arrow(d, (x1, y), (x2, y), width=2)
        bbox = d.textbbox((0, 0), text, font=_font(20))
        d.text(((x1 + x2) / 2 - (bbox[2] - bbox[0]) / 2, y - 32), text, fill="#334155", font=_font(20))
        y += 72
    path.parent.mkdir(parents=True, exist_ok=True)
    img.save(path)


def add_diagram(doc: Document, image_path: Path, caption: str, width: float = 6.4) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(width))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].italic = True
    cap.paragraph_format.space_after = Pt(8)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color: str = BORDER, size: str = "8") -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_column_widths(table, widths: list[float]) -> None:
    table.autofit = False
    total_twips = int(sum(widths) * 1440)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total_twips))
    tbl_w.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(int(width * 1440)))
        grid.append(col)

    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            cell.width = Inches(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")


def style_table(table, header_fill=LIGHT_BLUE, first_col_center=False) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row_idx, row in enumerate(table.rows):
        for col_idx, cell in enumerate(row.cells):
            set_cell_border(cell)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if row_idx == 0:
                set_cell_shading(cell, header_fill)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.color.rgb = RGBColor.from_string(NAVY)
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif first_col_center and col_idx == 0:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_page_number(section) -> None:
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("MyKaki / ElderGo KL Technical Design v2.0 | Page ")
    run.font.size = Pt(8)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Aptos Display"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos Display")
        run.font.color.rgb = RGBColor.from_string(NAVY if level == 1 else BLUE)
    if level == 1:
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(8)
    else:
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)


def add_para(doc: Document, text: str, style: str | None = None, bold_prefix: str | None = None) -> None:
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.08
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        r1.bold = True
        r1.font.color.rgb = RGBColor.from_string(NAVY)
        p.add_run(text[len(bold_prefix):])
    else:
        p.add_run(text)


def add_bullets(doc: Document, items: list[str], level: int = 0) -> None:
    style = "List Bullet" if level == 0 else "List Bullet 2"
    for item in items:
        p = doc.add_paragraph(style=style)
        p.paragraph_format.space_after = Pt(3)
        p.add_run(item)


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(3)
        p.add_run(item)


def add_callout(doc: Document, title: str, body: str, fill: str = LIGHT_ORANGE) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.right_indent = Inches(0.25)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(NAVY)
    p.add_run(": " + body)
    doc.add_paragraph()


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float], header_fill=LIGHT_BLUE) -> None:
    for row_data in rows:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        for i, value in enumerate(row_data):
            if i:
                p.add_run("; ")
            label = p.add_run(f"{headers[i]}: ")
            label.bold = True
            label.font.color.rgb = RGBColor.from_string(NAVY)
            p.add_run(value)
    doc.add_paragraph()


def add_simple_diagram(doc: Document, title: str, nodes: list[tuple[str, str]], note: str) -> None:
    add_para(doc, title)
    for i, (heading, body) in enumerate(nodes, start=1):
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(heading)
        r.bold = True
        r.font.color.rgb = RGBColor.from_string(NAVY)
        p.add_run(": " + body.replace("\n", "; "))
    doc.add_paragraph(note).italic = True
    doc.add_paragraph()


def build_doc() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    DIAGRAM_DIR.mkdir(parents=True, exist_ok=True)
    architecture_png = DIAGRAM_DIR / "figure_1_system_architecture.png"
    route_flow_png = DIAGRAM_DIR / "figure_2_route_planning_flow.png"
    backend_flow_png = DIAGRAM_DIR / "figure_3_backend_route_processing.png"
    station_flow_png = DIAGRAM_DIR / "figure_4_station_discovery_flow.png"
    weather_flow_png = DIAGRAM_DIR / "figure_5_weather_advice_sequence.png"
    preference_flow_png = DIAGRAM_DIR / "figure_6_preference_sync_sequence.png"
    deployment_png = DIAGRAM_DIR / "figure_7_render_deployment.png"
    save_architecture_diagram(architecture_png)
    save_vertical_flow(
        route_flow_png,
        "Figure 2: Route Planning and Result Workflow",
        [
            "User opens ElderGo KL",
            "Select origin from Google suggestions",
            "Select destination from Google suggestions",
            "Choose Now / Morning / Afternoon / Evening",
            "Frontend sends POST /routes/recommend",
            "Backend fetches Google transit candidates",
            "Backend scores, annotates, and returns one route",
            "Frontend shows summary, steps, map, weather, save, and share",
        ],
    )
    save_backend_flow(backend_flow_png)
    save_vertical_flow(
        station_flow_png,
        "Figure 4: Station Search and Detail Workflow",
        [
            "User opens Stations page",
            "Frontend requests /locations/popular",
            "User searches keyword when needed",
            "Frontend requests /locations/search",
            "Backend deduplicates station variants",
            "User selects station card",
            "Frontend requests /locations/{location_id}",
            "Station Detail shows routes, accessibility, facilities, sources, and More Details",
        ],
    )
    save_sequence_flow(
        weather_flow_png,
        "Figure 5: Weather Advice Sequence",
        ["User", "Frontend", "Backend", "OpenWeather"],
        [
            (0, 1, "Open route result", False),
            (1, 2, "POST /weather/forecast", False),
            (2, 3, "Request forecast", False),
            (3, 2, "Return forecast data", True),
            (2, 1, "Return risk + senior advice", True),
            (1, 0, "Show advisory card", True),
        ],
    )
    save_sequence_flow(
        preference_flow_png,
        "Figure 6: Preference and UI Settings Sync",
        ["User", "Frontend", "Local Storage", "Backend"],
        [
            (0, 1, "Change language/font/preferences", False),
            (1, 2, "Write local cache", False),
            (1, 3, "PATCH user settings/preferences", False),
            (3, 1, "Return saved state", True),
            (2, 1, "Fallback on restore failure", True),
        ],
    )
    save_vertical_flow(
        deployment_png,
        "Figure 7: Render Deployment Topology",
        [
            "Git repository contains frontend, backend, render.yaml",
            "Render auto-deploys eldergo-kl-frontend static service",
            "Render auto-deploys eldergo-kl-api Python web service",
            "Backend reads PostgreSQL, Google, OpenWeather, and CORS environment variables",
            "Frontend calls VITE_API_BASE_URL and uses SPA rewrite to index.html",
        ],
    )
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    add_page_number(section)

    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
    styles["Normal"].font.size = Pt(10.5)
    styles["Normal"].font.color.rgb = RGBColor.from_string("243447")
    for name in ("List Bullet", "List Bullet 2", "List Number"):
        styles[name].font.name = "Aptos"
        styles[name].font.size = Pt(10)

    # Cover page
    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("MyKaki / ElderGo KL")
    r.bold = True
    r.font.size = Pt(30)
    r.font.color.rgb = RGBColor.from_string(NAVY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("System Architecture and Technical Design Specification v2.0")
    r.bold = True
    r.font.size = Pt(20)
    r.font.color.rgb = RGBColor.from_string(BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Iteration 2: Senior-Friendly Klang Valley Travel Assistance")
    r.font.size = Pt(13)
    r.font.color.rgb = RGBColor.from_string("52616F")

    doc.add_paragraph()
    add_callout(
        doc,
        "Current Dev Branch Baseline",
        "This version describes the current ElderGo KL dev branch: React/Vite frontend, FastAPI backend, PostgreSQL/PostGIS data model, Google Maps and OpenWeather integrations, Render deployment, anonymous settings/preferences, station accessibility discovery, route annotations, and AI guardrail readiness.",
        LIGHT_BLUE,
    )
    doc.add_page_break()

    add_heading(doc, "Version Control", 1)
    add_para(doc, "This section records the document version baseline and the implementation snapshot used for this specification.")
    add_table(
        doc,
        ["Version", "Date", "Basis", "Summary"],
        [["2.0", "2026-04-28", "Current dev branch", "Rewritten for Iteration 2 architecture and current implemented behavior."],
         ["1.0", "Source reference", "Original MyKaki DOCX", "Used as structural reference only; outdated Iteration 1 claims were replaced."]],
        [0.8, 1.15, 1.55, 3.1],
        LIGHT_GREEN,
    )

    add_heading(doc, "Glossary", 1)
    add_table(
        doc,
        ["Term", "Meaning in Iteration 2"],
        [
            ["Anonymous user", "A backend user record resolved from a browser device ID hash; used for settings and travel preferences without full login."],
            ["UI settings", "Language, font-size mode, and onboarding completion state stored locally and synced to backend when possible."],
            ["Travel preferences", "Accessibility First, Least Walk, and Fewest Transfers flags used by route scoring."],
            ["Accessibility annotation", "A per-route-step message and status derived from station matching or nearby accessibility point data."],
            ["Searchable locations", "PostGIS-backed station and accessibility-related location records exposed through the Locations API."],
            ["PostGIS", "PostgreSQL spatial extension used for station geometry, route step geometry, and accessibility point proximity checks."],
            ["Google Places", "Autocomplete, place detail, and station image/detail support."],
            ["Google Maps Directions", "Transit route candidate source for backend route recommendation."],
            ["OpenWeather", "Forecast provider for destination weather advice."],
            ["Render", "Deployment platform for the frontend static service and backend API service."],
        ],
        [1.45, 5.15],
        LIGHT_BLUE,
    )

    add_heading(doc, "Table Of Contents", 1)
    toc_items = [
        "1. Document Purpose",
        "2. Project Overview",
        "3. Scope and Boundaries",
        "4. High-Level System Architecture",
        "5. Core Business Workflows",
        "6. Logical Architecture",
        "7. Data Design",
        "8. Deployment Architecture",
        "9. Security and Reliability",
        "10. Future Iteration Roadmap",
        "11. Review Conclusion",
        "Appendix A. Lean KIT",
        "Appendix B. Project Governance Portfolio",
    ]
    add_bullets(doc, toc_items)
    doc.add_page_break()

    add_heading(doc, "1. Document Purpose", 1)
    add_para(
        doc,
        "This document specifies the system architecture and technical design of MyKaki / ElderGo KL Iteration 2. It reflects the current dev branch, not the older Iteration 1 prototype assumptions.",
    )
    add_para(
        doc,
        "Iteration 2 focuses on senior-friendly route planning, persistent UI settings, travel preferences, station accessibility discovery, route accessibility annotations, weather advice, static help content, and AI guardrail readiness.",
    )

    add_heading(doc, "2. Project Overview", 1)
    add_para(
        doc,
        "ElderGo KL is a travel assistance application for older adults in Klang Valley, Malaysia. The current product is implemented as a mobile-first single-page web application with a FastAPI backend and data services for public transport, station discovery, and accessibility context.",
    )
    add_para(doc, "The current implemented modules are:")
    add_bullets(
        doc,
        [
            "Planning: origin and destination entry through Google Places autocomplete.",
            "Plan Your Time: Now, Morning, Afternoon, and Evening travel-time selection.",
            "Route Result: one recommended route, text/map views, step cards, weather advice, save as PNG, and sharing.",
            "Stations and Station Detail: popular/searchable station discovery with accessibility status, routes, known facilities, sources, and Google Maps detail lookup.",
            "Preference: accessibility-first, least-walk, and fewest-transfer settings.",
            "Help: app reset, Use ElderGo guide, ticket guide, concession guide, and privacy information.",
            "AI sheet: frontend quick-question chat surface; backend guardrail endpoints are prepared but not wired into the sheet yet.",
        ],
    )
    add_callout(
        doc,
        "Iteration 2 Corrections From v1.0",
        "The current branch does not implement browser current-location origin detection, route alternatives display, AQI display, ticket counter/operating-hour station fields, or live frontend-to-backend AI chat.",
        LIGHT_ORANGE,
    )

    add_heading(doc, "3. Scope and Boundaries", 1)
    add_heading(doc, "3.1 Implemented in Iteration 2", 2)
    add_bullets(
        doc,
        [
            "React + Vite frontend with mobile-first pages and global app context.",
            "FastAPI backend with routers for health, users, places, locations, routes, weather, and AI.",
            "Anonymous user resolution plus local/remote UI settings and travel preferences.",
            "Google Places autocomplete and place/station detail support.",
            "Google Maps transit route recommendation through `/routes/recommend`.",
            "Single-route scoring based on accessibility-first, walking distance, transfers, and duration.",
            "PostgreSQL/PostGIS station, route, accessibility, and searchable-location schema.",
            "Weather forecast advice via `/weather/forecast` and OpenWeather.",
            "Render deployment blueprint for frontend and backend services.",
        ],
    )
    add_heading(doc, "3.2 Not Included Yet", 2)
    add_bullets(
        doc,
        [
            "Full login, account management, and authenticated user profiles.",
            "True offline/PWA route storage; current route saving is a downloadable browser-generated PNG.",
            "Frontend chat integration with backend AI guardrail endpoints.",
            "AQI display or AQI endpoint.",
            "Ticket counter and operating-hour fields in station detail.",
            "Full real-time train arrival monitoring or real-time disruption handling.",
        ],
    )

    add_heading(doc, "4. High-Level System Architecture", 1)
    add_para(
        doc,
        "The solution follows a separated frontend-backend architecture with local browser fallback and external service orchestration. The frontend is responsible for senior-friendly interaction and state continuity; the backend validates requests, integrates data providers, annotates route accessibility, and persists data when configured.",
    )
    add_diagram(doc, architecture_png, "Figure 1: System Component Interaction Diagram (Iteration 2)", 6.6)
    add_para(doc, "The browser loads the static frontend, calls the FastAPI backend for application data, and uses local storage as a continuity fallback when backend restore fails.")

    add_heading(doc, "5. Core Business Workflows", 1)
    add_heading(doc, "5.1 Route Planning Flow", 2)
    add_diagram(doc, route_flow_png, "Figure 2: Route Planning and Result Workflow", 4.8)
    add_para(doc, "The app blocks unconfirmed text input and only continues once origin and destination have been selected from recognized suggestions.")
    add_heading(doc, "5.2 Backend Route Processing Flow", 2)
    add_diagram(doc, backend_flow_png, "Figure 3: Backend Route Request Processing Logic", 6.4)
    add_para(doc, "The backend returns a single `RecommendedRoute` object instead of exposing all candidate alternatives to the frontend.")
    add_heading(doc, "5.3 Station Discovery Flow", 2)
    add_diagram(doc, station_flow_png, "Figure 4: Station Search and Detail Workflow", 4.8)
    add_para(doc, "Station data is returned from PostGIS-backed searchable locations with CSV fallback for selected detail cases.")
    add_heading(doc, "5.4 Weather Advice Flow", 2)
    add_diagram(doc, weather_flow_png, "Figure 5: Weather Advice Sequence", 6.6)
    add_para(
        doc,
        "When a route result exists, the frontend requests `/weather/forecast` using destination name, coordinates when available, and the selected departure time. The backend calls OpenWeather, selects the nearest forecast period, calculates risk level, and returns senior-friendly advice. Weather is advisory only and does not replace the recommended route.",
    )
    add_heading(doc, "5.5 Preference and UI Settings Flow", 2)
    add_diagram(doc, preference_flow_png, "Figure 6: Preference and UI Settings Sync", 6.6)
    add_para(
        doc,
        "On app startup, the frontend creates or resolves an anonymous user from a generated browser device ID, restores UI settings and preferences when possible, and falls back to local storage if backend restore fails. Later changes are stored locally and patched back to the backend when an anonymous user ID exists.",
    )

    add_heading(doc, "6. Logical Architecture", 1)
    add_table(
        doc,
        ["Layer", "Current Components", "Responsibility"],
        [
            ["Presentation", "React pages, AppProvider, TopBar, BottomNav, translations", "Senior-friendly UI, navigation, local state, font/language switching, route/station/help flows."],
            ["API", "FastAPI routers for users, places, routes, locations, weather, AI, health", "Request validation, response contracts, backend boundary for frontend services."],
            ["Business services", "route_service, route_scoring_service, google_maps_service, station_matching_service, accessibility_annotation_service, weather_service, user_service", "Provider orchestration, route scoring, station matching, accessibility labeling, weather risk advice, preference persistence."],
            ["Data", "PostgreSQL/PostGIS schema, GTFS/static CSV import outputs", "Rail/station data, accessibility points, searchable locations, anonymous settings, route persistence, prepared AI tables."],
            ["External services", "Google Places, Google Maps Directions, OpenWeather", "Autocomplete, place/station detail, route candidates, forecast data."],
        ],
        [1.15, 2.25, 3.2],
        LIGHT_BLUE,
    )

    add_heading(doc, "7. Data Design", 1)
    add_para(
        doc,
        "Iteration 2 uses PostgreSQL/PostGIS as the current schema target. The schema is broader than the UI currently exposes, which allows route persistence, station discovery, accessibility annotations, and future saved-place or AI extensions without changing the core storage model.",
    )
    add_table(
        doc,
        ["Data Area", "Key Tables", "Purpose"],
        [
            ["Rail network", "rail_agencies, rail_routes, rail_stations, rail_station_routes", "Imported public transport structure and route membership."],
            ["Station grouping", "station_groups, station_group_members, station_accessibility_profiles", "Grouped station display names, cross-provider station membership, and accessibility summary."],
            ["Accessibility", "accessibility_points, route_accessibility_annotations", "Nearby accessibility features and per-step route accessibility messages."],
            ["Anonymous personalization", "anonymous_users, user_ui_settings, user_travel_preferences", "Device-based settings and route preference persistence without login."],
            ["Route planning", "route_requests, recommended_routes, route_steps", "Stores generated route requests, recommended route summaries, and ordered route instructions."],
            ["Discovery", "searchable_locations", "Station and location search with trigram and spatial indexes."],
            ["AI readiness", "ai_conversations, ai_messages", "Prepared storage model for future guarded assistant integration."],
        ],
        [1.35, 2.35, 2.9],
        LIGHT_GREEN,
    )
    add_heading(doc, "7.1 Public API Interfaces", 2)
    add_table(
        doc,
        ["Area", "Endpoint(s)", "Notes"],
        [
            ["Health", "`/`, `/health`, `/health/db`", "Root/API health and database connectivity checks."],
            ["Users", "`/users/anonymous`, `/users/{id}/ui-settings`, `/users/{id}/travel-preferences`", "Anonymous user resolution plus UI/preference restore and patch."],
            ["Places", "`/places/autocomplete`, `/places/details/{place_id}`, `/places/station-detail`, `/places/station-image`", "Google-backed place suggestions and detail/image support."],
            ["Routes", "`/routes/recommend`", "Single recommended route with steps and accessibility annotations."],
            ["Locations", "`/locations/popular`, `/locations/search`, `/locations/{location_id}`", "Popular station cards, keyword search, and station detail."],
            ["Weather", "`/weather/forecast`", "Destination forecast and senior-friendly risk advice."],
            ["AI", "`/ai/conversations`, `/ai/conversations/{conversation_id}/messages`", "Backend guardrail readiness; frontend sheet is not yet wired."],
        ],
        [1.05, 3.2, 2.35],
        LIGHT_BLUE,
    )
    add_heading(doc, "7.2 Frontend State Interfaces", 2)
    add_table(
        doc,
        ["Interface", "Fields", "Use"],
        [
            ["UISettings", "language, fontSize/font_size, onboardingCompleted/onboarding_completed", "Controls language, font-size mode, and first-time preference prompt state."],
            ["TravelPreferences", "accessibilityFirst/accessibility_first, leastWalk/least_walk, fewestTransfers/fewest_transfers", "Controls preference-aware route scoring."],
            ["RecommendedRoute", "route ID, origin/destination, duration, transfers, walking distance, reason, map polyline, steps, accessibility points", "Feeds Route Result summary, text cards, map embed, save, share, and weather context."],
        ],
        [1.45, 2.65, 2.5],
        LIGHT_ORANGE,
    )

    add_heading(doc, "8. Deployment Architecture", 1)
    add_para(
        doc,
        "The current deployment is defined in `render.yaml` as two Render services in one repository: a Python API service and a static frontend service.",
    )
    add_diagram(doc, deployment_png, "Figure 7: Render Deployment Topology", 4.9)
    add_para(doc, "The frontend calls `VITE_API_BASE_URL`, while backend secrets are injected through Render environment variables.")
    add_table(
        doc,
        ["Service", "Runtime", "Key Configuration"],
        [
            ["eldergo-kl-api", "Python web service", "Build: `pip install -r backend/requirements.txt`; Start: `gunicorn app.main:app --chdir backend -k uvicorn.workers.UvicornWorker`; health path `/api/v1/health` in Render config."],
            ["eldergo-kl-frontend", "Static service", "Build: `cd frontend && npm ci && npm run build`; publish path `frontend/dist`; SPA rewrite to `/index.html`."],
        ],
        [1.45, 1.45, 3.7],
        LIGHT_GREEN,
    )
    add_table(
        doc,
        ["Variable", "Scope", "Purpose"],
        [
            ["ELDERGO_ENV", "Backend", "Runtime environment label."],
            ["ELDERGO_DEMO_MODE", "Backend", "Allows demo/ephemeral behavior where persistence may be unavailable."],
            ["ELDERGO_CORS_ORIGINS", "Backend", "Allowed frontend origins."],
            ["ELDERGO_DATABASE_URL", "Backend secret", "PostgreSQL/PostGIS connection string."],
            ["ELDERGO_GOOGLE_MAPS_API_KEY", "Backend secret", "Google Directions and backend place support."],
            ["OPENWEATHER_API_KEY", "Backend secret", "OpenWeather forecast integration."],
            ["VITE_API_BASE_URL", "Frontend", "Backend API base URL."],
            ["VITE_GOOGLE_MAPS_API_KEY", "Frontend secret", "Browser-side map/station image integrations."],
        ],
        [2.05, 1.45, 3.1],
        LIGHT_BLUE,
    )

    add_heading(doc, "9. Security and Reliability", 1)
    add_bullets(
        doc,
        [
            "CORS is restricted through `ELDERGO_CORS_ORIGINS`, with local development origins configured by default.",
            "Google, OpenWeather, and database credentials are environment variables, not hard-coded in source.",
            "Anonymous user IDs avoid full account workflows; browser device IDs are represented through backend anonymous user resolution.",
            "UI settings and preferences are stored locally first and synchronized remotely when possible, keeping the prototype usable when backend restore fails.",
            "Health endpoints support API and database monitoring.",
            "Pydantic schemas validate route, weather, user, place, and AI request/response contracts.",
            "Location endpoints return explicit 503 responses when database-backed location data is unavailable.",
            "Route persistence falls back to demo or ephemeral route IDs so users can still receive a route result when database persistence is unavailable.",
        ],
    )

    add_heading(doc, "10. Future Iteration Roadmap", 1)
    add_numbered(
        doc,
        [
            "Add PWA/offline support for core pages, static assets, and saved route images/history.",
            "Wire the frontend AI sheet to backend guardrail endpoints and persist guarded conversations when appropriate.",
            "Add richer station fields only when reliable source data is available, such as operating hours, counter availability, and lift/escalator detail.",
            "Add AQI only after selecting a reliable API or dataset and defining risk advice rules for older adults.",
            "Improve accessibility-aware scoring with verified station/facility confidence rather than using walking/transfers as a proxy.",
            "Expose saved places or recent places through the existing schema once privacy and retention behavior is finalized.",
        ],
    )

    add_heading(doc, "11. Review Conclusion", 1)
    add_para(
        doc,
        "Iteration 2 has moved the product from a core route-planning prototype into a fuller elderly-friendly travel assistant. The current architecture supports route planning, preference-aware scoring, station accessibility discovery, weather advice, static guide content, and deployment-ready service separation.",
    )
    add_para(
        doc,
        "The branch is also honest about its boundaries: route alternatives, AQI, full AI integration, login, real-time transit data, and richer station operational fields remain future work rather than current behavior.",
    )

    doc.add_section(WD_SECTION.NEW_PAGE)
    add_heading(doc, "Appendix A. Lean KIT", 1)
    add_para(
        doc,
        "The Iteration 2 architecture supports lean delivery by separating user-facing experience, backend orchestration, and data enrichment. This keeps the prototype demonstrable while allowing future improvements to be added incrementally.",
    )
    add_table(
        doc,
        ["Lean Area", "Iteration 2 Position"],
        [
            ["Build", "Single-page frontend and FastAPI backend are independently deployable through Render."],
            ["Measure", "Backend tests cover route annotation persistence and location search/detail contracts."],
            ["Learn", "Current boundaries are explicit, so the next iteration can focus on high-value gaps such as PWA, AI wiring, and richer accessibility scoring."],
        ],
        [1.5, 5.1],
        LIGHT_GREEN,
    )

    add_heading(doc, "Appendix B. Project Governance Portfolio", 1)
    add_para(
        doc,
        "Governance for this version should treat the current dev branch as the technical source of truth. Any future document revision should re-check endpoints, schemas, deployment configuration, and implemented frontend pages before updating architecture claims.",
    )
    add_table(
        doc,
        ["Governance Item", "Current Control"],
        [
            ["Branch baseline", "Current `dev` branch."],
            ["Deployment baseline", "`render.yaml` services `eldergo-kl-api` and `eldergo-kl-frontend`."],
            ["Data baseline", "`backend/database/schema.sql` with PostgreSQL/PostGIS tables and indexes."],
            ["Verification baseline", "Backend tests under `backend/tests` plus visual DOCX render QA for this document."],
        ],
        [1.75, 4.85],
        LIGHT_BLUE,
    )

    doc.save(OUT_FILE)
    print(OUT_FILE)


if __name__ == "__main__":
    build_doc()
