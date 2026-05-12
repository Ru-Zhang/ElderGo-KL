from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "doc" / "generated"
DIAGRAM_DIR = OUT_DIR / "diagrams"
OUT_FILE = OUT_DIR / "MyKaki_v2_Route_Recommendation_Decision_Logic_Insert.docx"
DIAGRAM_FILE = DIAGRAM_DIR / "figure_route_recommendation_decision_logic.png"

NAVY = "1E3A5F"
BLUE = "4A90E2"
LIGHT_BLUE = "EAF3FE"
LIGHT_PURPLE = "F1ECFF"
PURPLE = "A88CFF"
LIGHT_GREEN = "ECF7EA"
LIGHT_ORANGE = "FFF3E8"
LIGHT_GRAY = "F5F7FA"
BORDER = "D9E2EC"
WHITE = "FFFFFF"


def _font(size: int, bold: bool = False):
    candidates = [
        Path("C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf"),
        Path("C:/Windows/Fonts/calibrib.ttf" if bold else "C:/Windows/Fonts/calibri.ttf"),
    ]
    for path in candidates:
        if path.exists():
            return ImageFont.truetype(str(path), size)
    return ImageFont.load_default()


def _wrap(draw: ImageDraw.ImageDraw, text: str, font, width: int) -> list[str]:
    words = text.split()
    lines: list[str] = []
    current = ""
    for word in words:
        candidate = f"{current} {word}".strip()
        if draw.textbbox((0, 0), candidate, font=font)[2] <= width:
            current = candidate
        else:
            if current:
                lines.append(current)
            current = word
    if current:
        lines.append(current)
    return lines


def _draw_box(draw, xy, text, fill="#F1ECFF", outline="#A88CFF", font_size=19, title=False):
    x1, y1, x2, y2 = xy
    if not fill.startswith("#"):
        fill = f"#{fill}"
    if not outline.startswith("#"):
        outline = f"#{outline}"
    draw.rounded_rectangle(xy, radius=10, fill=fill, outline=outline, width=3)
    font = _font(font_size, bold=title)
    lines = _wrap(draw, text, font, max(20, x2 - x1 - 28))
    line_h = font_size + 7
    total_h = len(lines) * line_h
    y = y1 + (y2 - y1 - total_h) / 2
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=font)
        draw.text((x1 + (x2 - x1 - (bbox[2] - bbox[0])) / 2, y), line, fill="#1E3A5F", font=font)
        y += line_h


def _draw_diamond(draw, center, size, text):
    cx, cy = center
    sx, sy = size
    points = [(cx, cy - sy), (cx + sx, cy), (cx, cy + sy), (cx - sx, cy)]
    draw.polygon(points, fill="#FFF3E8", outline="#E67E22")
    draw.line(points + [points[0]], fill="#E67E22", width=3)
    font = _font(19, bold=True)
    lines = _wrap(draw, text, font, sx * 2 - 34)
    line_h = 26
    y = cy - (len(lines) * line_h) / 2
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=font)
        draw.text((cx - (bbox[2] - bbox[0]) / 2, y), line, fill="#1E3A5F", font=font)
        y += line_h


def _arrow(draw, start, end, color="#334155", width=3):
    draw.line([start, end], fill=color, width=width)
    x1, y1 = start
    x2, y2 = end
    if abs(y2 - y1) >= abs(x2 - x1):
        direction = 1 if y2 >= y1 else -1
        pts = [(x2, y2), (x2 - 9, y2 - direction * 15), (x2 + 9, y2 - direction * 15)]
    else:
        direction = 1 if x2 >= x1 else -1
        pts = [(x2, y2), (x2 - direction * 15, y2 - 9), (x2 - direction * 15, y2 + 9)]
    draw.polygon(pts, fill=color)


def _label(draw, text, xy):
    draw.text(xy, text, fill="#64748B", font=_font(17, True))


def save_decision_logic_diagram(path: Path) -> None:
    img = Image.new("RGB", (1500, 1150), "white")
    d = ImageDraw.Draw(img)
    d.text((70, 40), "Figure 4: Route Recommendation Decision Logic", fill="#111827", font=_font(36, True))

    left_x = 90
    mid_x = 560
    right_x = 1040
    box_w = 360
    box_h = 86

    _draw_box(d, (left_x, 130, left_x + box_w, 130 + box_h), "User selects origin and destination from Google Places", fill=LIGHT_BLUE)
    _draw_box(d, (left_x, 270, left_x + box_w, 270 + box_h), "User selects schedule time: Now, Morning, Afternoon, Evening", fill=LIGHT_BLUE)
    _draw_box(d, (left_x, 410, left_x + box_w, 410 + box_h), "Frontend reads three travel preferences", fill=LIGHT_BLUE)
    _draw_box(d, (left_x, 550, left_x + box_w, 550 + box_h), "POST /routes/recommend with route inputs and preferences", fill=LIGHT_BLUE)

    _draw_box(d, (mid_x, 130, mid_x + box_w, 130 + box_h), "Backend validates origin, destination, time, and preferences", fill=LIGHT_PURPLE)
    _draw_box(d, (mid_x, 270, mid_x + box_w, 270 + box_h), "Request Google Directions transit candidates", fill=LIGHT_PURPLE)
    _draw_box(d, (mid_x, 410, mid_x + box_w, 410 + box_h), "Normalise candidates: duration, walking, transfers, steps", fill=LIGHT_PURPLE)
    _draw_box(d, (mid_x, 550, mid_x + box_w, 550 + box_h), "Apply preference-based route cost scoring", fill=LIGHT_PURPLE)
    _draw_diamond(d, (mid_x + box_w // 2, 750), (155, 95), "Best candidate found?")

    _draw_box(d, (right_x, 240, right_x + box_w, 240 + box_h), "No: return no-route or upstream provider error", fill=LIGHT_ORANGE, outline="#E67E22")
    _draw_box(d, (right_x, 550, right_x + box_w, 550 + box_h), "Yes: annotate transit and walking steps with local accessibility data", fill=LIGHT_GREEN, outline="#6BBF59")
    _draw_box(d, (right_x, 710, right_x + box_w, 710 + box_h), "Return one RecommendedRoute to frontend", fill=LIGHT_GREEN, outline="#6BBF59")
    _draw_box(d, (right_x, 870, right_x + box_w, 870 + box_h), "Frontend requests destination weather advice", fill=LIGHT_ORANGE, outline="#E67E22")
    _draw_box(d, (right_x, 1010, right_x + box_w, 1010 + box_h), "Show weather as advisory guidance only", fill=LIGHT_ORANGE, outline="#E67E22")

    for y in [216, 356, 496]:
        _arrow(d, (left_x + box_w // 2, y), (left_x + box_w // 2, y + 54))
    _arrow(d, (left_x + box_w, 593), (mid_x, 173))
    for y in [216, 356, 496]:
        _arrow(d, (mid_x + box_w // 2, y), (mid_x + box_w // 2, y + 54))
    _arrow(d, (mid_x + box_w // 2, 636), (mid_x + box_w // 2, 655))
    _arrow(d, (mid_x + box_w + 10, 750), (right_x, 593))
    _label(d, "Yes", (958, 652))
    _arrow(d, (mid_x + box_w // 2, 655), (right_x, 283))
    _label(d, "No", (898, 465))
    _arrow(d, (right_x + box_w // 2, 636), (right_x + box_w // 2, 710))
    _arrow(d, (right_x + box_w // 2, 796), (right_x + box_w // 2, 870))
    _arrow(d, (right_x + box_w // 2, 956), (right_x + box_w // 2, 1010))

    _label(d, "Recommendation factors", (170, 92))
    _label(d, "Backend scoring", (665, 92))
    _label(d, "Result and advisory layer", (1125, 202))

    path.parent.mkdir(parents=True, exist_ok=True)
    img.save(path)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_border(cell, color: str = BORDER) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=110, start=120, bottom=110, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def style_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.12

    for name, size, color in [
        ("Title", 22, NAVY),
        ("Heading 1", 16, NAVY),
        ("Heading 2", 13, BLUE),
        ("Heading 3", 11.5, NAVY),
    ]:
        style = styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)


def add_colored_paragraph(doc: Document, text: str, fill: str = LIGHT_BLUE) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_border(cell, BORDER)
    set_cell_margins(cell, 140, 160, 140, 160)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = RGBColor.from_string(NAVY)


def add_bullet(doc: Document, text: str, bold_prefix: str | None = None) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    if bold_prefix and text.startswith(bold_prefix):
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text[len(bold_prefix):])
    else:
        p.add_run(text)


def add_summary_table(doc: Document) -> None:
    rows = [
        ("Origin and destination", "Required route inputs", "Must be selected from Google Places suggestions before continuing."),
        ("Schedule time", "Participates in route request", "`now` is passed into Google Directions as current departure time. Other labels are captured for planning flow and future expansion."),
        ("Accessibility First", "Affects route scoring", "Prioritises lower walking distance and fewer transfers as practical accessibility proxies."),
        ("Least Walk", "Affects route scoring", "Increases the penalty for walking distance."),
        ("Fewest Transfers", "Affects route scoring", "Increases the penalty for route transfers."),
        ("Weather", "Advisory after recommendation", "Displayed on Route Result page; does not automatically change the recommended route in the current branch."),
        ("Accessibility annotations", "Added after candidate selection", "Google hints and local PostgreSQL/PostGIS station/accessibility data are used to explain route steps."),
    ]
    for factor, role, notes in rows:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        r1 = p.add_run(f"{factor}: ")
        r1.bold = True
        r1.font.color.rgb = RGBColor.from_string(NAVY)
        r2 = p.add_run(f"{role}. ")
        r2.bold = True
        p.add_run(notes.replace("`", ""))


def build_docx() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    save_decision_logic_diagram(DIAGRAM_FILE)

    doc = Document()
    style_doc(doc)

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("Route Recommendation Decision Logic Insert")

    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.LEFT
    note.paragraph_format.space_after = Pt(12)
    run = note.add_run("Where to paste: ")
    run.bold = True
    run.font.color.rgb = RGBColor.from_string(NAVY)
    note.add_run(
        "Place this section in Section 5: Core Business Workflows, immediately after 5.2 Backend Route Processing Flow. "
        "Use heading number 5.3, then renumber later Section 5 flows if needed."
    )

    doc.add_heading("5.3 Route Recommendation Decision Logic: Preferences, Schedule Time, and Weather Advice", level=1)
    doc.add_paragraph(
        "This workflow explains how ElderGo KL combines user route inputs, selected schedule time, and travel preferences "
        "to generate one senior-friendly route recommendation. The current Iteration 2 implementation separates route "
        "recommendation factors from post-recommendation advisory factors so the system remains transparent and testable."
    )
    doc.add_paragraph(
        "The route recommendation process begins when the user selects a valid origin and destination from Google Places "
        "suggestions and then chooses a travel time option. The frontend sends the selected places, selected departure time, "
        "anonymous user ID when available, and the three travel preferences to the backend /routes/recommend endpoint."
    )

    doc.add_heading("Travel Preferences Used in Scoring", level=2)
    add_bullet(
        doc,
        "Accessibility First: prioritises route candidates with lower walking burden and fewer transfers because these are currently the most reliable measurable proxies for elderly-friendly travel.",
        "Accessibility First:",
    )
    add_bullet(doc, "Least Walk: gives stronger priority to routes with shorter walking distance.", "Least Walk:")
    add_bullet(doc, "Fewest Transfers: gives stronger priority to routes with fewer transit changes.", "Fewest Transfers:")

    doc.add_heading("Schedule Time Handling", level=2)
    doc.add_paragraph(
        "The selected schedule time affects the route search by being passed into the Google Directions request. In the current "
        "implementation, the Now option is sent directly as departure_time=now, allowing Google Directions to calculate candidate "
        "routes based on current transit availability. Other time labels are captured by the frontend flow and kept in the route "
        "request model, but deeper scheduled departure handling can be expanded in a later iteration."
    )

    doc.add_heading("Preference-Based Route Cost", level=2)
    doc.add_paragraph(
        "After Google Directions returns candidate transit routes, ElderGo KL normalises each candidate into comparable values "
        "such as total duration, total walking distance, transfer count, route steps, and route polyline. The backend then "
        "calculates a route cost using the user preferences. Lower route cost is better. When a preference is enabled, that "
        "preference has stronger influence on the ordering:"
    )
    add_bullet(doc, "If Accessibility First is enabled, candidates with less walking and fewer transfers are considered first.")
    add_bullet(doc, "If Least Walk is enabled, walking distance becomes a stronger decision factor.")
    add_bullet(doc, "If Fewest Transfers is enabled, transfer count becomes a stronger decision factor.")
    add_bullet(doc, "Duration remains part of the weighted cost so the system avoids selecting an unreasonable route only because it has fewer transfers or less walking.")

    doc.add_paragraph(
        "The backend returns one recommended route rather than a list of alternatives. This design reduces decision burden for "
        "older adults and presents a single route with step-by-step guidance, accessibility annotations, weather advice, save "
        "options, and sharing options."
    )

    doc.add_page_break()
    doc.add_heading("Weather Advice Boundary", level=2)
    doc.add_paragraph(
        "Weather is currently handled as advisory information after route recommendation, not as an automatic route replacement "
        "factor. Once the recommended route is available, the frontend requests weather advice for the route destination and "
        "selected travel time context. The Route Result page displays senior-friendly weather guidance, such as reminders about "
        "rain, heat, or preparation, but it does not automatically recalculate or replace the recommended route in Iteration 2."
    )
    doc.add_paragraph(
        "This distinction is important for system reliability. Preferences and schedule time participate in the route recommendation "
        "flow, while weather currently helps the user make an informed decision after the recommendation is shown. A future iteration "
        "can extend this design by introducing weather-aware scoring, for example by increasing walking penalties during heavy rain "
        "or high heat, but this behaviour is not part of the current dev-branch implementation."
    )

    doc.add_page_break()
    doc.add_heading("Figure 4: Route Recommendation Decision Logic", level=2)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(DIAGRAM_FILE), width=Inches(6.55))
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = caption.add_run("Figure 4: Route Recommendation Decision Logic")
    run.italic = True
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor.from_string(NAVY)

    doc.add_heading("Current Iteration 2 Behaviour Summary", level=2)
    add_summary_table(doc)

    doc.add_page_break()
    doc.add_heading("Future Enhancement: Weather-Aware Scoring", level=2)
    doc.add_paragraph(
        "If Iteration 3 introduces weather-aware route scoring, the current design can be extended without changing the main "
        "route-planning user flow. The backend can request weather before final route selection and adjust the route cost when "
        "weather increases risk for older adults."
    )
    add_bullet(doc, "Increase walking-distance penalty during heavy rain.")
    add_bullet(doc, "Increase walking-distance penalty during high heat.")
    add_bullet(doc, "Prefer fewer transfers during poor weather to reduce outdoor waiting.")
    add_bullet(doc, "Show a stronger warning when no lower-risk route is available.")
    doc.add_paragraph(
        "This should be presented as a future enhancement unless the backend route scoring service is updated to use weather data "
        "before selecting the recommended route."
    )

    doc.save(OUT_FILE)


if __name__ == "__main__":
    build_docx()
